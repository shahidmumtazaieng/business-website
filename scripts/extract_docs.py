#!/usr/bin/env python3
"""Extract content from all strategic docx files for analysis."""
import os
from docx import Document

UPLOAD_DIR = "/home/z/my-project/upload"
OUTPUT_DIR = "/home/z/my-project/scripts/extracted"
os.makedirs(OUTPUT_DIR, exist_ok=True)

docx_files = [f for f in os.listdir(UPLOAD_DIR) if f.endswith(".docx")]
print(f"Found {len(docx_files)} docx files\n")

for fname in docx_files:
    fpath = os.path.join(UPLOAD_DIR, fname)
    print(f"=== Extracting: {fname} ===")
    try:
        doc = Document(fpath)
        out_lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                if "Heading" in style:
                    out_lines.append(f"\n## [{style}] {text}\n")
                else:
                    out_lines.append(text)

        # Also extract tables
        for ti, table in enumerate(doc.tables):
            out_lines.append(f"\n[TABLE {ti+1}]")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                out_lines.append(" | ".join(cells))
            out_lines.append("")

        content = "\n".join(out_lines)
        out_fname = fname.replace(".docx", ".txt").replace(" ", "_").replace(":", "")
        out_path = os.path.join(OUTPUT_DIR, out_fname)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"  -> Saved {len(content)} chars to {out_fname}")
    except Exception as e:
        print(f"  ERROR: {e}")
print("\nDone.")
